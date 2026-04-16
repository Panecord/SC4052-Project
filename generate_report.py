"""
Generate the CloneMe Final Report as a DOCX file.
Run: python generate_report.py
Output: CloneMe_Final_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size  = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        elif level == 2:
            run.font.size  = Pt(13)
            run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)
        else:
            run.font.size  = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_para(text="", bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # Shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F3F4F6")
    pPr.append(shd)
    return p

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.bold      = header

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_picture  # no cover image — use styled text

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("CloneMe")
run.font.name  = "Calibri"
run.font.size  = Pt(36)
run.bold       = True
run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("A Personal Assistant-as-a-Service Platform")
run2.font.name  = "Calibri"
run2.font.size  = Pt(18)
run2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

doc.add_paragraph()

for line in [
    "NTU SC4053 Cloud Computing — Final Project Report",
    "Topic 2: Personal Assistant-as-a-Service",
    "",
    "Author: Iman Dzafir",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "",
    "Source Code: https://github.com/[GITHUB_USERNAME]/CloneMe",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    if "NTU" in line or "Topic" in line:
        run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Abstract", level=1)
add_para(
    "This report presents CloneMe, a Personal Assistant-as-a-Service (PAaaS) platform built "
    "as part of the NTU SC4053 Cloud Computing course. CloneMe decomposes personal assistant "
    "capabilities into independent, loosely-coupled microservices that communicate via well-defined "
    "RESTful APIs. Powered by Anthropic's Claude AI as an agentic reasoning engine, CloneMe "
    "integrates Gmail, Google Calendar, GitHub, Mastodon, HackerNews, Reddit, and YouTube APIs to "
    "automate scheduling, email management, real-time trend analysis, and AI-driven application "
    "development. The platform demonstrates the full PAaaS lifecycle: a natural-language chat "
    "interface drives tool-use agents that read emails, create calendar events, analyse social trends, "
    "and autonomously generate and deploy GitHub repositories through a vibe-coding workflow. "
    "Server-Sent Events (SSE) provide real-time streaming feedback to the user. The system "
    "exemplifies the principles of service decomposition, API-driven development, and cloud-native "
    "architecture with a focus on practical utility and demonstrable originality."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", level=1)
toc_items = [
    ("1", "Introduction", "3"),
    ("2", "Literature Review", "4"),
    ("3", "Problem Statement", "6"),
    ("4", "Solution Design and Architecture", "7"),
    ("  4.1", "System Overview", "7"),
    ("  4.2", "Backend Microservices (Flask)", "8"),
    ("  4.3", "AI Agent Engine", "9"),
    ("  4.4", "Email Service", "10"),
    ("  4.5", "Calendar Service", "11"),
    ("  4.6", "Trends Service", "11"),
    ("  4.7", "GitHub / Vibe-Coding Service", "12"),
    ("  4.8", "Research Pipeline Tools", "13"),
    ("  4.9", "Build History Manager", "14"),
    ("  4.10", "Frontend Architecture", "14"),
    ("  4.11", "Security and Authentication", "15"),
    ("5", "Illustrative Examples", "15"),
    ("6", "Evaluation", "18"),
    ("7", "Conclusions and Future Work", "19"),
    ("8", "References", "20"),
]
t = doc.add_table(rows=0, cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(0.6)
t.columns[1].width = Inches(4.2)
t.columns[2].width = Inches(0.8)
for num, title, page in toc_items:
    row = t.add_row()
    for i, val in enumerate([num, title, page]):
        cell = row.cells[i]
        cell.text = val
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 2 else WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
# Remove table borders — just use spacing
from docx.oxml import OxmlElement as OE
tbl = t._tbl
tblPr = tbl.tblPr
tblBorders = OE("w:tblBorders")
for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
    b = OE(f"w:{border_name}")
    b.set(qn("w:val"),   "none")
    b.set(qn("w:sz"),    "0")
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), "auto")
    tblBorders.append(b)
tblPr.append(tblBorders)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Introduction", level=1)
add_para(
    "The rapid proliferation of cloud-based services, AI language models, and open APIs has created "
    "an unprecedented opportunity to build 'digital clones' — intelligent personal assistants that "
    "know a user's habits, manage their digital life, and autonomously execute complex multi-step "
    "tasks on their behalf. This vision is embodied by the fictional OpenClaw project referenced in "
    "the assignment brief, which achieved over 100,000 GitHub stars by offering capabilities ranging "
    "from reading email to booking flights, all while building a persistent model of user preferences."
)
add_para(
    "This project, CloneMe, is a working implementation of that vision. It demonstrates the "
    "Personal Assistant-as-a-Service (PAaaS) paradigm by decomposing assistant capabilities into "
    "independent, API-connected microservices orchestrated by a Claude AI reasoning engine. The "
    "platform operates through a natural-language chat interface that hides the complexity of "
    "multi-service coordination from the end user, presenting a seamless experience while internally "
    "dispatching structured tool calls across email, calendar, trend analysis, and code generation "
    "services."
)
add_para(
    "The key contributions of this work are:"
)
add_bullet("A working PAaaS platform with five integrated service domains (email, calendar, trends, code generation, task management)")
add_bullet("An agentic AI loop using Claude's tool-use API with Server-Sent Events for real-time streaming")
add_bullet("A vibe-coding workflow that autonomously generates, pushes, and previews GitHub repositories from natural language prompts")
add_bullet("A multi-source trend intelligence aggregator spanning Mastodon, HackerNews, GitHub, Reddit, and YouTube")
add_bullet("A build history manager that tracks, previews, and deletes AI-generated repositories")
add_bullet("OAuth2 with PKCE for Google services and IMAP support for Outlook and Yahoo Mail")

doc.add_paragraph()
add_heading("1.1 Scope", level=2)
add_para(
    "CloneMe is a single-user, locally-hosted application intended for personal use and academic "
    "demonstration. It runs as a Flask server (Python 3.11+) and is accessed via a modern web browser. "
    "All AI processing is handled by Anthropic's Claude API. Sensitive credentials (API keys, OAuth "
    "tokens) are stored locally in .env and JSON token files. The system is not multi-tenant by "
    "design, though the service decomposition pattern would permit multi-tenancy with appropriate "
    "authentication middleware."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Literature Review", level=1)

add_heading("2.1 Cloud Computing and Microservices", level=2)
add_para(
    "Cloud computing has evolved from infrastructure provisioning (IaaS) to platform services (PaaS) "
    "and software delivery (SaaS), with the latest frontier being AI-driven agent services [1]. "
    "The microservices architectural pattern, popularised by Netflix and Amazon, advocates decomposing "
    "monolithic applications into small, independently deployable services each owning a single "
    "business capability [2]. This decomposition improves fault isolation, enables independent "
    "scaling, and allows polyglot technology stacks. CloneMe adopts this philosophy at the application "
    "level: each assistant capability (email, calendar, trends, GitHub) is implemented as an "
    "independent service module with its own API endpoints."
)

add_heading("2.2 AI Agents and Tool Use", level=2)
add_para(
    "The introduction of tool-use (function calling) in large language models [3] fundamentally "
    "changed the role of AI from passive responders to active agents capable of taking actions in "
    "external systems. The ReAct pattern (Reasoning + Acting) [4] describes a feedback loop where "
    "the model alternates between generating reasoning traces and invoking tools, iterating until "
    "the user's goal is satisfied. Anthropic's Claude models implement this pattern natively through "
    "a structured tool-use API where tools are declared as JSON schemas and the model generates "
    "structured tool_call blocks that the application must execute and return as tool_result messages."
)
add_para(
    "OpenAI's GPT-4 function calling [5], AutoGPT [6], and LangChain [7] are related systems that "
    "demonstrate the power of agentic AI. CloneMe's agent loop is architecturally similar but "
    "implemented with a focus on streaming, real-time feedback, and a minimal dependency footprint."
)

add_heading("2.3 Personal Assistant Systems", level=2)
add_para(
    "Virtual assistants such as Apple Siri, Google Assistant, and Amazon Alexa established the "
    "consumer use case for AI-driven task delegation [8]. Academic work on personal assistant "
    "architectures has explored modular designs [9], contextual awareness [10], and habit modelling "
    "over time [11]. More recent systems such as Microsoft Copilot and Google Gemini Advanced "
    "integrate directly with productivity suites (email, calendar, documents), demonstrating the "
    "commercial viability of the PAaaS model. CloneMe differs from these proprietary systems by "
    "providing an open, extensible architecture that the user controls entirely."
)

add_heading("2.4 API-Driven Development (Vibe-Coding)", level=2)
add_para(
    "The term 'vibe-coding', coined in AI developer communities circa 2024-2025, refers to "
    "AI-assisted software development where the developer communicates intent in natural language "
    "and the AI generates, revises, and deploys code. Cursor IDE, GitHub Copilot Workspace, and "
    "Anthropic's own Claude Code tool exemplify this workflow. The GitHub API's Contents API enables "
    "programmatic creation and modification of repository files, making fully autonomous code "
    "generation pipelines feasible [12]. CloneMe's Build feature implements a complete vibe-coding "
    "loop: the user states what they want, the AI generates the full codebase, creates a private "
    "GitHub repository, pushes all files, saves a local preview, and records the build in a "
    "persistent registry."
)

add_heading("2.5 Social Trend Analysis via APIs", level=2)
add_para(
    "Real-time trend intelligence has significant applications in research, marketing, and product "
    "development. Mastodon's open ActivityPub-based federated network provides a trends API "
    "(/api/v1/trends/tags) not subject to the paywalls of X/Twitter [13]. HackerNews (Y Combinator) "
    "offers a public Firebase-based API for fetching stories, comments, and engagement data [14]. "
    "Reddit's JSON API (/.json suffix) provides community-curated content ranking [15]. GitHub's "
    "search API allows discovery of rapidly-starred repositories [16]. YouTube's Data API v3 "
    "surfaces trending videos [17]. CloneMe aggregates all five sources into a unified trend "
    "intelligence feed with per-source detail enrichment."
)

add_heading("2.6 OAuth2 and API Security", level=2)
add_para(
    "OAuth 2.0 is the industry standard for delegated API authorisation [18]. The Proof Key for "
    "Code Exchange (PKCE) extension [19] adds security to the authorization code flow for public "
    "clients by preventing authorization code interception attacks. Google's OAuth2 implementation "
    "requires PKCE for web applications, necessitating careful state management across HTTP redirects. "
    "IMAP (Internet Message Access Protocol) [20] remains the dominant standard for accessing email "
    "from third-party applications on providers such as Yahoo Mail, though Microsoft's Exchange "
    "Online has deprecated Basic Authentication in favour of OAuth2/Modern Authentication, creating "
    "compatibility challenges for third-party IMAP clients."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Problem Statement", level=1)
add_para(
    "Modern knowledge workers interact with a fragmented ecosystem of digital services: email in "
    "one application, calendar in another, social trends in several browser tabs, and development "
    "tools in yet another environment. Synthesising information across these silos is time-consuming "
    "and cognitively taxing. Existing solutions (Siri, Google Assistant) are either locked to "
    "proprietary ecosystems, lack programmability, or do not provide the kind of deep integration "
    "with development tools that a technically-oriented user requires."
)
add_para(
    "Specifically, the following unmet needs motivate this project:"
)
add_bullet(
    "Unified AI assistant: No open-source platform combines email management, calendar scheduling, "
    "trend monitoring, and AI code generation in a single locally-controlled application."
)
add_bullet(
    "Autonomous task execution: Existing tools respond to queries but rarely take multi-step actions "
    "autonomously (e.g., 'read my unread emails, identify the most important one, and schedule a "
    "follow-up meeting')."
)
add_bullet(
    "API-driven development automation: Converting a natural language idea into a deployed GitHub "
    "repository typically requires switching between multiple tools and platforms."
)
add_bullet(
    "Cross-platform trend intelligence: Aggregating trends from Mastodon, HackerNews, GitHub, "
    "Reddit, and YouTube requires five separate browser tabs and manual synthesis."
)
add_bullet(
    "Build accountability: AI-generated code repositories accumulate without a management interface; "
    "users have no easy way to review, access, or delete builds."
)
add_para(
    "CloneMe addresses all five of these needs within a single, locally-hosted platform."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SOLUTION DESIGN
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Solution Design and Architecture", level=1)

add_heading("4.1 System Overview", level=2)
add_para(
    "CloneMe is a single-process Flask application that logically decomposes into six independent "
    "service domains. The architecture follows a hub-and-spoke model: the AI agent engine is the "
    "hub, and each service (email, calendar, trends, GitHub, tasks) is a spoke that the agent "
    "invokes via well-defined internal tool functions. External integrations are accessed via their "
    "respective REST APIs over HTTPS."
)

# Architecture table
add_para("Table 1: System Component Overview", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
t = doc.add_table(rows=1, cols=4)
t.style = "Light Shading Accent 1"
headers = ["Component", "Technology", "External API", "Purpose"]
for i, h in enumerate(headers):
    cell = t.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

rows_data = [
    ["AI Agent Engine",       "Python + Anthropic SDK",      "Claude API (claude-sonnet-4-6)",        "Reasoning, tool orchestration, SSE streaming"],
    ["Email Service",         "Python + imaplib + Gmail API","Gmail API v1 / IMAP (Outlook, Yahoo)",  "Read, search, summarise emails"],
    ["Calendar Service",      "Python + Google API Client",  "Google Calendar API v3",                "Fetch events, create appointments"],
    ["Trends Service",        "Python + requests",           "Mastodon, HN, GitHub, Reddit, YouTube", "Aggregate real-time trend data"],
    ["GitHub Service",        "Python + requests",           "GitHub REST API v3",                    "Create repos, push files, star repos, vibe-coding"],
    ["Research Pipeline",     "Python + ddgs + requests",    "DuckDuckGo Search, any URL",            "Web search + webpage reader to inform builds"],
    ["Task Manager",          "Python threading + dict",     "—",                                     "Background job queue with status tracking"],
    ["Build Registry",        "JSON file + Python",          "GitHub REST API v3 (delete)",           "Persist, display, delete AI-generated builds"],
    ["Frontend",              "Vanilla JS + CSS + HTML",     "—",                                     "Single-page app, SSE consumer, multi-tab"],
]
for row_data in rows_data:
    row = t.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)

doc.add_paragraph()

add_heading("4.2 Backend Microservices (Flask)", level=2)
add_para(
    "The backend is a single Flask application (app.py, approximately 1,700 lines) that exposes "
    "RESTful API endpoints for each service domain. Despite being a monolith at the process level, "
    "the code is architecturally decomposed into service modules, each with its own route namespace, "
    "helper functions, and external API interactions. This design choice trades the operational "
    "complexity of a true microservices deployment (Kubernetes, service discovery, inter-service "
    "authentication) for practical simplicity while preserving the architectural clarity of service "
    "separation."
)
add_para("Key API endpoint namespaces:")
add_bullet("/api/chat             — POST, streams AI agent response via SSE")
add_bullet("/api/clarify          — POST, intent classification before routing")
add_bullet("/api/gmail/*          — Gmail inbox, message fetch, pagination")
add_bullet("/api/email/*          — IMAP provider management (connect, inbox, message)")
add_bullet("/api/calendar/*       — Events fetch, event creation, OAuth callback")
add_bullet("/api/trends, /api/trends/detail — Multi-source trend aggregation + per-item detail")
add_bullet("/api/builds, /api/builds/<name> — Build registry CRUD + GitHub delete")
add_bullet("/api/tasks            — Background task status polling")
add_bullet("/preview/<path>       — Static file server for AI-generated app previews")
add_bullet("/api/google/auth, /api/google/callback — Google OAuth2 PKCE flow")

doc.add_paragraph()
add_para(
    "Flask-CORS is enabled to facilitate local development. The Flask session is used with a "
    "configurable secret key (FLASK_SECRET_KEY) for OAuth state management."
)

add_heading("4.3 AI Agent Engine", level=2)
add_para(
    "The AI agent is built on Anthropic's claude-sonnet-4-6 model using the tool-use (function "
    "calling) capability. The agent implements a standard ReAct loop:"
)
add_bullet("The user sends a natural-language prompt via POST /api/chat")
add_bullet("The server calls the Anthropic messages.stream() API with a system prompt, conversation history, and tool definitions")
add_bullet("The model streams text tokens and tool_call blocks as Server-Sent Events (SSE) to the browser")
add_bullet("When a tool_call is received, the server executes the corresponding Python function and appends the result as a tool_result message")
add_bullet("The loop continues until the model emits a stop_reason of 'end_turn' with no pending tool calls")

add_para("The SSE event stream uses the following event types:")
add_code_block("data: {\"type\": \"status\",      \"message\": \"My clone is working…\"}")
add_code_block("data: {\"type\": \"text\",         \"content\": \"Based on your inbox…\"}")
add_code_block("data: {\"type\": \"tool_call\",    \"id\": \"tc_001\", \"name\": \"get_email_inbox\", \"input\": {...}}")
add_code_block("data: {\"type\": \"tool_result\",  \"id\": \"tc_001\", \"name\": \"get_email_inbox\", \"result\": {...}}")
add_code_block("data: {\"type\": \"done\"}")

add_para(
    "The frontend distinguishes between 'info tools' (email, calendar, trends) that display an "
    "animated spinner status message, and 'build tools' (create_github_repo, push_files_to_repo) "
    "that display expandable tool cards showing inputs and outputs. This design provides appropriate "
    "visual feedback without exposing low-level JSON to the user for routine read operations."
)
add_para(
    "Three Claude model tiers are offered: claude-haiku-4-5-20251001 (fastest, cheapest), "
    "claude-sonnet-4-6 (default, balanced), and claude-opus-4-6 (most capable). The user selects "
    "the model per-session from a dropdown in the chat interface."
)

add_heading("4.4 Email Service", level=2)
add_para(
    "The email service supports three providers:"
)
add_bullet("Gmail — via Google Gmail API v1 with OAuth2 authentication. Supports search queries (Gmail search syntax), pagination via nextPageToken, and full message body fetch with HTML-to-text conversion.")
add_bullet("Outlook — via IMAP over SSL (imap-mail.outlook.com:993). Note: Microsoft has permanently disabled Basic Authentication for Exchange Online, so @outlook.com personal accounts are supported but corporate/school accounts on Exchange Online require OAuth2/Modern Authentication (not yet implemented).")
add_bullet("Yahoo Mail — via IMAP over SSL (imap.mail.yahoo.com:993) using Yahoo App Passwords.")

add_para(
    "The Gmail route defaults to 30 emails per page with 'Load more' pagination. The email list "
    "renders in a two-pane layout: a scrollable message list on the left and a reading pane on the "
    "right. The AI agent can search, read, and summarise emails; draft replies; and answer questions "
    "about inbox content. A friendly error message is displayed when Exchange Online's Basic Auth "
    "block is detected."
)

add_heading("4.5 Calendar Service", level=2)
add_para(
    "The calendar service integrates with Google Calendar API v3 through the same OAuth2 credentials "
    "used for Gmail (single-connection flow). Features include:"
)
add_bullet("7-day week grid view with Monday-start ISO weeks, today highlight, and weekend shading")
add_bullet("Week navigation (previous/next/today) with localised date display")
add_bullet("All-day and timed event support with timezone-safe rendering (local date parts, avoiding UTC/SGT offset issues)")
add_bullet("Event detail modal showing title, time range, location, and description")
add_bullet("New Event creation form with datetime pickers")
add_bullet("Real-time calendar refresh when the AI agent creates an event (SSE-triggered refresh with 400ms debounce)")

add_para(
    "The AI agent's create_calendar_event tool accepts ISO 8601 datetimes and the Asia/Singapore "
    "timezone by default, making it natural to say 'schedule a meeting for tomorrow at 2pm'."
)

add_heading("4.6 Trends Service", level=2)
add_para(
    "The trends service aggregates content from five platforms and provides two API endpoints:"
)
add_bullet("/api/trends?source=X — returns a list of trend items for the selected source (mastodon, hackernews, github, youtube, reddit, or all)")
add_bullet("/api/trends/detail?source=X&url=Y&tag=Z — fetches rich detail for a single trend item")

add_para("Per-source implementations:")
add_bullet("Mastodon: Trending hashtags from /api/v1/trends/tags and trending statuses from /api/v1/trends/statuses. Detail view shows recent posts for a hashtag with boost and like counts.")
add_bullet("HackerNews: Top Show HN stories from the Firebase API with score and comment counts. Detail view shows the story metadata and top 5 comments.")
add_bullet("GitHub: Recently-starred Python/JavaScript/TypeScript repositories from the Search API. Detail view shows star count, fork count, language, topics, and README excerpt.")
add_bullet("Reddit: Top posts from r/programming via Reddit's JSON API. Detail view shows post body and top comments.")
add_bullet("YouTube: Trending videos from YouTube Data API v3 (chart=mostPopular, regionCode=SG). Detail view shows thumbnail, view count, like count, channel name, and description.")

add_para(
    "The trends grid uses CSS auto-fill with minmax(280px, 1fr), making it fully responsive. "
    "Clicking any trend card opens a detail modal with source-specific rich content and one-click "
    "Analyse (sends to AI chat) or Build (generates an app) actions."
)

add_heading("4.7 GitHub / Vibe-Coding Service", level=2)
add_para(
    "The GitHub service implements a complete API-driven development cycle. When the user asks "
    "CloneMe to 'build' something, the AI agent executes a two-step workflow:"
)
add_bullet("Step 1 — create_github_repo: Calls POST /user/repos with private: true and auto_init: true to create a new private repository.")
add_bullet("Step 2 — push_files_to_repo: Calls PUT /repos/{owner}/{repo}/contents/{path} for each file in the generated codebase. Existing files are updated by fetching their SHA first. Files are also written to a local previews/ directory for instant browser preview.")

add_para("Example vibe-coding prompt and response flow:")
add_code_block("User: Build a Pomodoro timer web app with dark mode")
add_code_block("")
add_code_block("Agent creates: pomodoro-timer (private GitHub repo)")
add_code_block("Agent pushes:  index.html, style.css, app.js, README.md")
add_code_block("Preview URL:   http://localhost:5000/preview/pomodoro-timer/")
add_code_block("GitHub URL:    https://github.com/{user}/pomodoro-timer")

add_para(
    "The AI generates complete, self-contained web applications. All generated HTML, CSS, and "
    "JavaScript is written in a single batch call to minimise API round-trips. The Claude model "
    "is instructed to write clean, modern, functional code without external CDN dependencies "
    "wherever possible."
)
add_para(
    "An optional post_to_mastodon tool allows the agent to announce newly built apps to the user's "
    "Mastodon feed, completing the social sharing aspect of the vibe-coding workflow."
)

add_heading("4.8 Research Pipeline Tools", level=2)
add_para(
    "Three tools were added specifically to strengthen the connection between the Trends and "
    "GitHub services, forming a research-then-build pipeline:"
)
add_bullet(
    "web_search — Queries DuckDuckGo (via the ddgs package, no API key required) and returns "
    "titles, URLs, and snippets. The agent uses this to research a trending topic before building: "
    "'What are the best features of a Pomodoro timer?' → search → incorporate findings into the "
    "generated app."
)
add_bullet(
    "read_webpage — Fetches any URL and returns its readable text content (scripts and style blocks "
    "stripped, HTML entities decoded). The agent uses this to read the actual HackerNews article, "
    "GitHub README, Reddit post, or YouTube video description behind a trend card's source_url, "
    "giving it the context needed to generate a much more accurate and relevant app."
)
add_bullet(
    "star_github_repo — Calls PUT /user/starred/{owner}/{repo} on the GitHub API to star a "
    "repository on the user's behalf. This directly connects the Trends feed (GitHub trending "
    "cards) to the GitHub service: the user can say 'star this repo for me' and the agent acts "
    "on it immediately."
)
add_para(
    "Together, these tools create a coherent research loop: discover a trend → read the source "
    "material → search for related best practices → build an informed app → star the repos that "
    "inspired it. All three appear as animated spinner status messages in the chat UI (no JSON "
    "cards) to keep the experience clean."
)

add_heading("4.9 Build History Manager", level=2)
add_para(
    "The build history manager addresses the problem of AI-generated repository accumulation. "
    "All builds are persisted in a builds.json registry file with the following schema:"
)
add_code_block("{")
add_code_block("  \"repo_name\":   \"pomodoro-timer\",")
add_code_block("  \"description\": \"A Pomodoro timer web app with dark mode\",")
add_code_block("  \"repo_url\":    \"https://github.com/user/pomodoro-timer\",")
add_code_block("  \"full_name\":   \"user/pomodoro-timer\",")
add_code_block("  \"created_at\":  \"2026-04-16T14:23:00+00:00\",")
add_code_block("  \"status\":      \"ready\",")
add_code_block("  \"files_count\": 4,")
add_code_block("  \"pushed_at\":   \"2026-04-16T14:23:45+00:00\",")
add_code_block("  \"preview_url\": \"/preview/pomodoro-timer/\"")
add_code_block("}")

add_para(
    "The Build History modal (accessible via the chat interface) displays all builds as cards "
    "showing the repo name, description, creation date, file count, and status badge. Each card "
    "provides Preview (local), GitHub ↗ (open repo), and Delete buttons. Delete calls "
    "DELETE /api/builds/<repo_name>, which removes the GitHub repository via the API, deletes "
    "the local preview directory, and removes the registry entry — a complete cleanup in one click."
)

add_heading("4.10 Frontend Architecture", level=2)
add_para(
    "The frontend is a single-page application (SPA) built with vanilla JavaScript (no framework), "
    "modular CSS custom properties, and semantic HTML5. The JavaScript is split into seven modules "
    "loaded in dependency order:"
)
add_bullet("state.js — shared mutable state (active tab index, current trend source, email provider)")
add_bullet("utils.js — XSS-safe escaper (esc()), relative time formatter (relativeTime()), markdown renderer")
add_bullet("chat.js — tab system, SSE consumer, message feed renderer, tool card builder, build flow")
add_bullet("email.js — provider tabs, inbox list, email reader pane, IMAP modal")
add_bullet("calendar.js — week grid renderer, event detail modal, new event form")
add_bullet("trends.js — trend cards, detail modal, per-source rich content renderer")
add_bullet("tasks.js — background task list, build history modal, delete flow")
add_bullet("app.js — navigation, section switching, global keyboard shortcuts, health polling")

add_para(
    "All user-visible strings are sanitised through the esc() function before insertion into "
    "innerHTML, preventing XSS vulnerabilities from API data. The UI uses CSS Grid and Flexbox "
    "with CSS custom properties (design tokens) for a consistent visual language. The layout is "
    "fully responsive with adaptive padding for screens below 900px."
)

add_heading("4.11 Security and Authentication", level=2)
add_para(
    "Security considerations addressed in CloneMe:"
)
add_bullet("Google OAuth2 with PKCE: The code_verifier is written to a file-based fallback (.oauth_pkce.tmp) in addition to the Flask session, guarding against session cookie loss during the OAuth redirect. The file is deleted immediately after use.")
add_bullet("API key management: All credentials (Anthropic, GitHub, Google, Mastodon, YouTube) are stored in a .env file and never exposed to the frontend.")
add_bullet("XSS prevention: All data rendered into HTML is passed through an esc() function that escapes &, <, >, \", and ' characters.")
add_bullet("Path traversal prevention: The DELETE /api/builds/<name> endpoint validates the repo_name parameter against a whitelist regex (^[a-zA-Z0-9_\\-.]+$) before using it in file system operations.")
add_bullet("Input validation: Calendar event datetimes are validated with ISO 8601 parsing; IMAP credentials are only stored locally and never sent to the AI model.")
add_bullet("HTTPS: All external API calls use HTTPS. The local Flask server uses HTTP for localhost development only.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ILLUSTRATIVE EXAMPLES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Illustrative Examples", level=1)
add_para(
    "This section demonstrates CloneMe's capabilities through five end-to-end usage scenarios."
)

add_heading("Example 1: Email Intelligence and Action", level=2)
add_para("Scenario: The user asks CloneMe to summarise their unread emails and identify any requiring action.")
add_code_block("User: What unread emails do I have and what needs my attention?")
add_code_block("")
add_code_block("[AI spins 'Reading emails…' indicator]")
add_code_block("→ tool_call: get_email_inbox({\"query\": \"is:unread\", \"max_results\": 20})")
add_code_block("← tool_result: {emails: [{subject: 'Meeting reschedule request', from: 'boss@company.com'}, ...")
add_code_block("")
add_code_block("Agent: You have 8 unread emails. Most urgent:")
add_code_block("  1. 'Meeting reschedule request' from your boss — sent 2 hours ago,")
add_code_block("     requesting to move Thursday's 3pm meeting to Friday.")
add_code_block("  2. 'Invoice #2341 overdue' from finance@vendor.com — payment deadline tomorrow.")
add_code_block("  3. 4 newsletters (Substack, GitHub Digest) — can be reviewed at your leisure.")
add_code_block("")
add_code_block("Would you like me to draft a reply to the meeting reschedule or")
add_code_block("create a calendar reminder for the invoice payment?")
add_para(
    "The AI agent autonomously decides to use the get_email_inbox tool, processes the results, "
    "and presents a prioritised summary without the user needing to open their email client."
)

add_heading("Example 2: Calendar Management", level=2)
add_para("Scenario: The user asks the AI to schedule a meeting based on context from their emails.")
add_code_block("User: Schedule a meeting with my boss tomorrow at 2pm about the reschedule request")
add_code_block("")
add_code_block("[AI spins 'Fetching calendar…' indicator]")
add_code_block("→ tool_call: get_calendar_events({\"days_ahead\": 2})")
add_code_block("← tool_result: {events: [{summary: 'Team standup', start: '2026-04-17T09:00:00'}]}")
add_code_block("")
add_code_block("[AI spins 'Creating calendar event…' indicator]")
add_code_block("→ tool_call: create_calendar_event({")
add_code_block("    \"summary\": \"Meeting with Boss — Schedule Discussion\",")
add_code_block("    \"start_datetime\": \"2026-04-17T14:00:00\",")
add_code_block("    \"end_datetime\":   \"2026-04-17T15:00:00\",")
add_code_block("    \"timezone\": \"Asia/Singapore\"")
add_code_block("  })")
add_code_block("← tool_result: {success: true, link: 'https://calendar.google.com/...'}")
add_code_block("")
add_code_block("Agent: Done! I've created 'Meeting with Boss — Schedule Discussion'")
add_code_block("for tomorrow (17 Apr) at 2:00–3:00 PM SGT. The calendar view has")
add_code_block("been updated.")
add_para(
    "The calendar grid in the UI automatically refreshes 400ms after the event creation tool result "
    "is received, showing the new event without any manual page reload."
)

add_heading("Example 3: Trend Analysis and App Building", level=2)
add_para("Scenario: The user discovers a trending topic and asks CloneMe to build a related app.")
add_code_block("User: [clicks 'Build →' on the '#Pomodoro' HackerNews trending card]")
add_code_block("")
add_code_block("Prompt auto-filled: 'Build a web app inspired by: Building a Focus Timer")
add_code_block("with Gamification — Show HN submission with 847 points'")
add_code_block("")
add_code_block("→ tool_call: create_github_repo({")
add_code_block("    \"name\": \"focus-timer-gamified\",")
add_code_block("    \"description\": \"Gamified Pomodoro focus timer inspired by HN\"")
add_code_block("  })")
add_code_block("← tool_result: {success: true, repo_url: 'https://github.com/user/focus-timer-gamified'}")
add_code_block("")
add_code_block("→ tool_call: push_files_to_repo({")
add_code_block("    \"repo_name\": \"focus-timer-gamified\",")
add_code_block("    \"files\": [")
add_code_block("      {\"path\": \"index.html\", \"content\": \"<!DOCTYPE html>...\"},")
add_code_block("      {\"path\": \"style.css\",  \"content\": \":root { --purple: #6B21A8... }\"},")
add_code_block("      {\"path\": \"app.js\",     \"content\": \"const WORK_MINS = 25...\"},")
add_code_block("      {\"path\": \"README.md\",  \"content\": \"# Focus Timer...\"}")
add_code_block("    ]")
add_code_block("  })")
add_code_block("← tool_result: {success: true, files_pushed: 4, preview_url: '/preview/focus-timer-gamified/'}")
add_code_block("")
add_code_block("Agent: Built! Your gamified focus timer is live:")
add_code_block("  • Preview: http://localhost:5000/preview/focus-timer-gamified/")
add_code_block("  • GitHub:  https://github.com/user/focus-timer-gamified (private)")
add_code_block("  4 files pushed. Recorded in Build History.")

add_heading("Example 4: Cross-Platform Trend Intelligence", level=2)
add_para("Scenario: The user clicks 'Analyse All Trends' to get an AI briefing on what's trending.")
add_code_block("User: Fetch trends from Mastodon, HackerNews, and Reddit, then give me a brief")
add_code_block("analysis: what are the top 5 themes, most interesting, and what should I know?")
add_code_block("")
add_code_block("→ tool_call: fetch_trends({\"source\": \"mastodon\"})")
add_code_block("← tool_result: {items: [{title: '#AIArt', ...}, {title: '#ClimateAction', ...}, ...]}")
add_code_block("")
add_code_block("→ tool_call: fetch_trends({\"source\": \"hackernews\"})")
add_code_block("← tool_result: {items: [{title: 'Building with Claude Tool Use', ...}, ...]}")
add_code_block("")
add_code_block("→ tool_call: fetch_trends({\"source\": \"reddit\"})")
add_code_block("← tool_result: {items: [{title: 'Rust 2026 edition preview', ...}, ...]}")
add_code_block("")
add_code_block("Agent: Top themes across platforms today:")
add_code_block("  1. AI/LLM tooling (HN dominant — 3 of top 10 stories)")
add_code_block("  2. Climate & sustainability (Mastodon #ClimateAction trending)")
add_code_block("  3. Rust language adoption (Reddit r/programming, r/rust)")
add_code_block("  4. Open-source AI art tools (#AIArt, #StableDiffusion Mastodon)")
add_code_block("  5. WebAssembly in production (HN Show HN story with 600+ points)")
add_para(
    "The AI synthesises across three separate API calls, demonstrating multi-source reasoning "
    "that would require manual effort to replicate."
)

add_heading("Example 5: Build History Management", level=2)
add_para("Scenario: The user reviews their builds and deletes one they no longer need.")
add_code_block("User: [opens Build History modal]")
add_code_block("")
add_code_block("Build History shows:")
add_code_block("  • focus-timer-gamified   [ready]  Apr 16 2026 · 4 files  [Preview] [GitHub↗] [Delete]")
add_code_block("  • weather-dashboard      [ready]  Apr 15 2026 · 6 files  [Preview] [GitHub↗] [Delete]")
add_code_block("  • portfolio-site         [ready]  Apr 10 2026 · 5 files  [Preview] [GitHub↗] [Delete]")
add_code_block("")
add_code_block("User: [clicks Delete on 'weather-dashboard']")
add_code_block("Confirm: Delete 'weather-dashboard'?")
add_code_block("  • Delete the GitHub repo permanently")
add_code_block("  • Remove the local preview files")
add_code_block("")
add_code_block("User: [confirms]")
add_code_block("→ DELETE /api/builds/weather-dashboard")
add_code_block("  ├─ GitHub API: DELETE /repos/user/weather-dashboard → 204 No Content")
add_code_block("  ├─ Local:      shutil.rmtree(previews/weather-dashboard)")
add_code_block("  └─ Registry:   removed from builds.json")
add_code_block("← {success: true, github_deleted: true}")
add_code_block("")
add_code_block("Build list refreshes: weather-dashboard is gone from GitHub and locally.")

doc.add_page_break()

add_heading("Example 6: Research-Driven App Build", level=2)
add_para(
    "Scenario: The user sees a trending GitHub repo in the trends feed and asks CloneMe to "
    "read it, research best practices, then build a similar app — demonstrating the full "
    "research pipeline from trend discovery to deployed code."
)
add_code_block("User: [clicks 'Analyse' on GitHub trending card: 'Excalidraw — Virtual whiteboard']")
add_code_block("")
add_code_block("Prompt: Analyse the trend: Build a web app similar to Excalidraw:")
add_code_block("         An open source virtual hand-drawn style whiteboard")
add_code_block("")
add_code_block("[AI spins 'Reading page…']")
add_code_block("→ tool_call: read_webpage({\"url\": \"https://github.com/excalidraw/excalidraw\"})")
add_code_block("← tool_result: {text: \"Excalidraw is a virtual collaborative whiteboard tool...\"}")
add_code_block("               {text: \"...supports hand-drawn style, real-time collaboration,")
add_code_block("                       local-first, export to PNG/SVG, dark mode...\"}")
add_code_block("")
add_code_block("[AI spins 'Searching the web…']")
add_code_block("→ tool_call: web_search({\"query\": \"canvas drawing app HTML5 best features 2025\"})")
add_code_block("← tool_result: {results: [")
add_code_block("    {title: 'Building a Canvas Drawing App — best practices', snippet: '...'},")
add_code_block("    {title: 'Hand-drawn style with rough.js library', snippet: '...'},")
add_code_block("  ]}")
add_code_block("")
add_code_block("[AI spins 'Starring repo…']")
add_code_block("→ tool_call: star_github_repo({\"repo_full_name\": \"excalidraw/excalidraw\"})")
add_code_block("← tool_result: {success: true}")
add_code_block("")
add_code_block("→ tool_call: create_github_repo({\"name\": \"sketch-canvas\",")
add_code_block("    \"description\": \"Hand-drawn style HTML5 canvas whiteboard\"})")
add_code_block("→ tool_call: push_files_to_repo({\"repo_name\": \"sketch-canvas\",")
add_code_block("    \"files\": [index.html (Canvas API + rough.js CDN), style.css, app.js,")
add_code_block("               README.md with feature list from research]})")
add_code_block("")
add_code_block("Agent: Built 'sketch-canvas' — a hand-drawn style whiteboard with:")
add_code_block("  • Freehand drawing, shapes (rectangle, circle, line)")
add_code_block("  • Rough.js sketch aesthetic (discovered via web search)")
add_code_block("  • Dark/light mode toggle, PNG export")
add_code_block("  • Preview: http://localhost:5000/preview/sketch-canvas/")
add_code_block("  • GitHub: https://github.com/user/sketch-canvas (private)")
add_code_block("  I've also starred excalidraw/excalidraw for your reference.")
add_para(
    "This example illustrates the complete research pipeline: the AI reads the actual GitHub "
    "README (via read_webpage) to understand the project's scope, searches for implementation "
    "best practices (via web_search), stars the original repo for later reference "
    "(via star_github_repo), and builds an informed app that incorporates findings from both "
    "research steps — all from a single click on a trend card."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Evaluation", level=1)

add_heading("6.1 Feature Completeness", level=2)
add_para("Table 2: Feature Implementation Summary", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
t2 = doc.add_table(rows=1, cols=3)
t2.style = "Light Shading Accent 2"
for i, h in enumerate(["Feature", "Status", "Notes"]):
    t2.rows[0].cells[i].text = h
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t2.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

features = [
    ["Natural language chat interface",          "✓ Implemented", "Multi-tab, SSE streaming"],
    ["Gmail integration (OAuth2)",               "✓ Implemented", "Read, search, summarise, paginate"],
    ["Outlook IMAP integration",                 "✓ Implemented", "Personal accounts; Exchange blocked"],
    ["Yahoo Mail IMAP integration",              "✓ Implemented", "App Password required"],
    ["Google Calendar read/write",               "✓ Implemented", "7-day view, event creation, real-time refresh"],
    ["AI-driven calendar event creation",        "✓ Implemented", "Natural language → ISO datetime"],
    ["Mastodon trend aggregation",               "✓ Implemented", "Tags + statuses with detail modal"],
    ["HackerNews trend aggregation",             "✓ Implemented", "Show HN + comments"],
    ["GitHub trending repos",                    "✓ Implemented", "Stars-sorted, topics, README"],
    ["Reddit trend aggregation",                 "✓ Implemented", "r/programming JSON API"],
    ["YouTube trending videos",                  "✓ Implemented", "Data API v3, SG region"],
    ["Vibe-coding (GitHub repo generation)",     "✓ Implemented", "Create + push + preview"],
    ["Web search (research pipeline)",           "✓ Implemented", "DuckDuckGo, no API key, informs builds"],
    ["Webpage reader (research pipeline)",       "✓ Implemented", "Fetch + extract any URL before building"],
    ["Star GitHub repo from trends",             "✓ Implemented", "GitHub API, one-click from trend card"],
    ["Build history manager",                    "✓ Implemented", "Persist, view, delete"],
    ["Mastodon posting",                         "✓ Implemented", "Status post with hashtags"],
    ["Background task tracking",                 "✓ Implemented", "UUID tasks, live badge polling"],
    ["Multi-model selection",                    "✓ Implemented", "Haiku / Sonnet / Opus"],
    ["Responsive layout (all sections)",         "✓ Implemented", "CSS Grid, no max-width constraint"],
    ["AI loading animation",                     "✓ Implemented", "Pulsing purple banner + spinner"],
    ["PKCE OAuth2 fix",                          "✓ Implemented", "File-based code_verifier fallback"],
]
for row_data in features:
    row = t2.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            if row_data[1] == "✓ Implemented" and i == 1:
                run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph()

add_heading("6.2 API Integration Summary", level=2)
add_para(
    "CloneMe integrates with nine distinct external APIs: Anthropic Claude API, Google Gmail API v1, "
    "Google Calendar API v3, GitHub REST API v3, Mastodon API v1, HackerNews Firebase API, Reddit "
    "JSON API, YouTube Data API v3, and IMAP (standard protocol for Outlook/Yahoo). This breadth "
    "of integration demonstrates the PAaaS decomposition principle: each domain is handled by its "
    "most appropriate provider API."
)

add_heading("6.3 Limitations", level=2)
add_bullet("Single-user: no multi-tenancy, no user accounts, no database — by design for this scope")
add_bullet("Exchange Online: Microsoft's deprecated Basic Auth for Exchange Online prevents @e.ntu.edu.sg IMAP access; OAuth2 Modern Auth is the correct fix but was not implemented in this project")
add_bullet("No persistent conversation history: each chat session starts fresh; long-term user preference modelling (as described for OpenClaw) is future work")
add_bullet("Gmail write: the current Gmail scope is read-only; sending emails requires additional OAuth scopes")
add_bullet("Local preview only: AI-generated apps are served from localhost; cloud deployment (e.g., Vercel, GitHub Pages) would require additional infrastructure")
add_bullet("Rate limits: YouTube Data API has a 10,000 unit/day quota; heavy trend usage could exhaust it")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Conclusions and Future Work", level=1)

add_heading("7.1 Conclusions", level=2)
add_para(
    "CloneMe demonstrates that the Personal Assistant-as-a-Service paradigm is achievable with "
    "a small codebase (~2,500 lines of Python + ~2,500 lines of JavaScript) and commodity cloud "
    "APIs. The key architectural insight is that a powerful AI agent engine, combined with "
    "well-designed tool definitions, can coordinate multiple independent service APIs to accomplish "
    "complex multi-step tasks that would previously require bespoke integration work."
)
add_para(
    "The vibe-coding workflow is particularly novel: a user can describe a web application in "
    "natural language, and CloneMe will generate the full codebase, create a GitHub repository, "
    "push all files, provide a live local preview, and record the build in a managed registry — "
    "all in under 60 seconds. This demonstrates the API-driven development cycle described in "
    "the assignment brief."
)
add_para(
    "The multi-source trend intelligence aggregator, real-time SSE streaming, responsive UI, and "
    "OAuth2/PKCE security implementation each represent engineering challenges solved in the course "
    "of this project, collectively delivering a platform that is both practically useful and "
    "architecturally sound as a demonstration of cloud computing principles."
)

add_heading("7.2 Future Work", level=2)
add_bullet("Exchange Online / Microsoft Graph API integration for NTU and corporate email accounts")
add_bullet("Persistent user preference modelling: store conversation history and extracted preferences in SQLite or a vector database for truly personalised responses over time")
add_bullet("Gmail write access: send emails, create draft replies, archive/label messages via the AI agent")
add_bullet("Cloud deployment: containerise with Docker, deploy to AWS/GCP with proper secret management (AWS Secrets Manager, GCP Secret Manager)")
add_bullet("Multi-tenancy: add user accounts with per-user credential storage and isolated conversation contexts")
add_bullet("Voice interface: integrate Web Speech API for voice-to-text input and text-to-speech output")
add_bullet("Proactive notifications: schedule periodic AI-driven briefings (e.g., morning email summary pushed to Mastodon or a notification service)")
add_bullet("Plugin architecture: allow third parties to contribute new service modules via a defined tool interface")
add_bullet("Mobile-responsive layout: refine for phone-screen usage with collapsible navigation")
add_bullet("GitHub Pages deployment: automatically deploy AI-generated apps to GitHub Pages for public access")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. References", level=1)

references = [
    "[1]  Amazon Web Services. (2024). \"What is Cloud Computing?\" Amazon Web Services. https://aws.amazon.com/what-is-cloud-computing/",
    "[2]  Fowler, M., & Lewis, J. (2014). \"Microservices.\" martinfowler.com. https://martinfowler.com/articles/microservices.html",
    "[3]  Anthropic. (2024). \"Tool use (function calling).\" Anthropic Documentation. https://docs.anthropic.com/en/docs/build-with-claude/tool-use",
    "[4]  Yao, S., et al. (2022). \"ReAct: Synergizing Reasoning and Acting in Language Models.\" arXiv:2210.03629.",
    "[5]  OpenAI. (2023). \"Function calling and other API updates.\" OpenAI Blog. https://openai.com/index/function-calling-and-other-api-updates/",
    "[6]  Significant Gravitas. (2023). \"AutoGPT: An Autonomous GPT-4 Experiment.\" GitHub. https://github.com/Significant-Gravitas/AutoGPT",
    "[7]  LangChain AI. (2023). \"LangChain Documentation.\" https://docs.langchain.com/",
    "[8]  Hoy, M. B. (2018). \"Alexa, Siri, Cortana, and More: An Introduction to Voice Assistants.\" Medical Reference Services Quarterly, 37(1), 81-88.",
    "[9]  Maes, P. (1994). \"Agents that reduce work and information overload.\" Communications of the ACM, 37(7), 30-40.",
    "[10] Abowd, G. D., et al. (1999). \"Towards a Better Understanding of Context and Context-Awareness.\" HUC '99.",
    "[11] Lashkari, A. H., et al. (1994). \"Collaborative Interface Agents.\" AAAI 1994.",
    "[12] GitHub. (2024). \"REST API — Repositories contents.\" GitHub Docs. https://docs.github.com/en/rest/repos/contents",
    "[13] Mastodon. (2024). \"REST API — Trends.\" Mastodon Documentation. https://docs.joinmastodon.org/methods/trends/",
    "[14] HackerNews. (2024). \"The Hacker News API.\" GitHub. https://github.com/HackerNews/API",
    "[15] Reddit Inc. (2024). \"Reddit API Documentation.\" https://www.reddit.com/dev/api/",
    "[16] GitHub. (2024). \"REST API — Search repositories.\" GitHub Docs. https://docs.github.com/en/rest/search/search#search-repositories",
    "[17] Google. (2024). \"YouTube Data API v3 Reference.\" Google Developers. https://developers.google.com/youtube/v3/docs/videos/list",
    "[18] Hardt, D. (2012). \"The OAuth 2.0 Authorization Framework.\" RFC 6749. IETF.",
    "[19] Sakimura, N., et al. (2015). \"Proof Key for Code Exchange by OAuth Public Clients.\" RFC 7636. IETF.",
    "[20] Crispin, M. (2003). \"INTERNET MESSAGE ACCESS PROTOCOL — VERSION 4rev1.\" RFC 3501. IETF.",
    "[21] Anthropic. (2025). \"Claude Model Overview.\" Anthropic Documentation. https://docs.anthropic.com/en/docs/about-claude/models/overview",
    "[22] Flask. (2024). \"Flask Documentation (3.x).\" https://flask.palletsprojects.com/",
    "[23] Google Identity. (2024). \"OAuth 2.0 for Web Server Applications.\" Google Developers. https://developers.google.com/identity/protocols/oauth2/web-server",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = "Calibri"
    run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("Appendix A: Environment Setup", level=1)
add_para("To run CloneMe locally:", size=11)
add_code_block("# 1. Clone the repository")
add_code_block("git clone https://github.com/[GITHUB_USERNAME]/CloneMe.git")
add_code_block("cd CloneMe")
add_code_block("")
add_code_block("# 2. Create virtual environment")
add_code_block("python -m venv .venv")
add_code_block(".venv\\Scripts\\activate   # Windows")
add_code_block("")
add_code_block("# 3. Install dependencies")
add_code_block("pip install -r requirements.txt")
add_code_block("")
add_code_block("# 4. Configure environment variables")
add_code_block("cp .env.example .env")
add_code_block("# Edit .env and fill in:")
add_code_block("#   ANTHROPIC_API_KEY=sk-ant-...")
add_code_block("#   GITHUB_TOKEN=ghp_...")
add_code_block("#   GITHUB_USERNAME=your_username")
add_code_block("#   MASTODON_ACCESS_TOKEN=...")
add_code_block("#   MASTODON_API_BASE_URL=https://mastodon.social")
add_code_block("#   YOUTUBE_API_KEY=...")
add_code_block("#   GOOGLE_CLIENT_ID=...")
add_code_block("#   GOOGLE_CLIENT_SECRET=...")
add_code_block("#   FLASK_SECRET_KEY=your-secret-key")
add_code_block("")
add_code_block("# 5. Run the server")
add_code_block("python app.py")
add_code_block("")
add_code_block("# 6. Open in browser")
add_code_block("http://localhost:5000")

add_heading("Appendix B: Tool Definitions Summary", level=1)
add_para("The AI agent has access to the following tools:", size=11)
tools_table_data = [
    ("create_github_repo",    "name, description",                    "Create a new private GitHub repository"),
    ("push_files_to_repo",    "repo_name, files[], commit_message",   "Push source files to GitHub and save locally for preview"),
    ("post_to_mastodon",      "status",                               "Post a status update to Mastodon (max 500 chars)"),
    ("get_email_inbox",       "max_results, query",                   "Fetch Gmail inbox with optional search query"),
    ("get_calendar_events",   "days_ahead, from_date",                "Fetch upcoming Google Calendar events"),
    ("create_calendar_event", "summary, start_datetime, end_datetime, description, timezone", "Create a new calendar event"),
    ("fetch_trends",          "source (mastodon|hackernews|github|youtube|reddit|all)", "Fetch trending content from the specified platform"),
    ("web_search",            "query, max_results",                   "Search the web via DuckDuckGo to research topics before building"),
    ("read_webpage",          "url, max_chars",                       "Fetch and extract readable text from any URL (HN articles, READMEs, etc.)"),
    ("star_github_repo",      "repo_full_name (owner/repo)",          "Star a GitHub repository on the user's behalf"),
]
t3 = doc.add_table(rows=1, cols=3)
t3.style = "Light Shading"
for i, h in enumerate(["Tool Name", "Parameters", "Description"]):
    t3.rows[0].cells[i].text = h
    t3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t3.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
for name, params, desc in tools_table_data:
    row = t3.add_row()
    for i, val in enumerate([name, params, desc]):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            run.font.name = "Calibri" if i > 0 else "Courier New"
            run.font.size = Pt(9.5)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "CloneMe_Final_Report_v2.docx"
doc.save(output_path)
print(f"Report saved: {output_path}")
